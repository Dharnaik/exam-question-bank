import io
from datetime import datetime
import re
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

VERSION = "KWv3-2025-08-10 (MIN+py311)"

# ---------------- minimal RAKE-style extractor (no spaCy) ----------------
_STOP = set("""
the and or for with from into onto about using use uses of in on to by as at this that which these those their your our its are is be been being it an a do does did can could may might should would will shall how what why when where
define description describe explain discuss briefly write list state name give show draw calculate solve determine find compute compare analyze analyse assess evaluate identify demonstrate apply design develop construct create formulate justify argue critique outline summarize classify distinguish examine
question answer marks unit subunit co teacher id year frequency keywords blooms taxonomy course outcome tag type
""".split())
_GENERIC = {"introduction","overview","system","method","methods","process","processes","concept","types","factors","steps",
            "advantages","disadvantages","merits","demerits","importance","role","effect","effects","impact","principles",
            "principle","purpose","function","functions","components","parameters","features","example","examples"}
_WORD = re.compile(r"[A-Za-z0-9][A-Za-z0-9\-]*")

def _strip_edge_stops(tokens):
    i, j = 0, len(tokens) - 1
    while i <= j and (tokens[i].lower() in _STOP or len(tokens[i]) <= 1): i += 1
    while j >= i and (tokens[j].lower() in _STOP or len(tokens[j]) <= 1): j -= 1
    return tokens[i:j+1]

def _clean_phrase(tokens):
    toks = _strip_edge_stops(tokens)
    if not toks or len(toks) < 2: return ""
    text = " ".join(toks).lower()
    if text in _GENERIC: return ""
    return text

def _candidates_rake(text: str):
    out, seen = [], set()
    for q in re.findall(r"['\"]([^'\"]{3,80})['\"]", text):
        ph = _clean_phrase([t for t in _WORD.findall(q)])
        if ph and ph not in seen: seen.add(ph); out.append(ph)
    tokens = _WORD.findall(text)
    chunks, cur = [], []
    for w in tokens:
        lw = w.lower()
        if lw in _STOP:
            if cur: chunks.append(cur); cur=[]
        else: cur.append(w)
    if cur: chunks.append(cur)
    freq, deg = {}, {}
    for ch in chunks:
        for w in ch:
            lw = w.lower()
            freq[lw] = freq.get(lw, 0) + 1
            deg[lw]  = deg.get(lw, 0) + (len(ch) - 1)
    scores = {w: (deg[w] + freq[w]) / freq[w] for w in freq}
    scored = []
    for ch in chunks:
        if 2 <= len(ch) <= 6:
            ph = _clean_phrase([w.lower() for w in ch])
            if not ph or ph in seen: continue
            s = sum(scores.get(w.lower(), 1.0) for w in ch)
            scored.append((s, ph))
    scored.sort(key=lambda x: (-x[0], -len(x[1])))
    for _, ph in scored:
        if ph not in seen: seen.add(ph); out.append(ph)
    return out[:3] if out else ["NO_PHRASE_FOUND"]

# ---------------- DOCX helpers ----------------
def _set_cell(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    f = run.font
    f.name = "Calibri"
    f.size = Pt(11)

def _row(table, label, value, bold_val=False):
    c = table.add_row().cells
    _set_cell(c[0], label)
    _set_cell(c[1], value, bold_val)

def detect_bloom_level(question):
    q = str(question).lower()
    bloom = {
        "L1": ["define","list","name","state","identify","recall"],
        "L2": ["explain","describe","summarize","classify","outline"],
        "L3": ["solve","use","demonstrate","compute","apply"],
        "L4": ["compare","differentiate","analyze","distinguish","examine"],
        "L5": ["justify","evaluate","assess","argue","critique"],
        "L6": ["design","develop","formulate","construct","create"],
    }
    for lvl, verbs in bloom.items():
        for v in verbs:
            if f" {v} " in f" {q} ": return lvl
    return "L2"

def assign_difficulty(bloom_level):
    return {"L1":"Low","L2":"Low","L3":"Medium","L4":"Medium","L5":"High","L6":"High"}.get(bloom_level,"Medium")

def classify_question_type(question):
    return "P" if any(w in str(question).lower() for w in ["calculate","solve","determine","find","compute"]) else "T"

# ---------------- Streamlit UI ----------------
st.title("📚 Question Bank — Minimal (RAKE only) + DOCX")
st.caption(f"App version: {VERSION}")

qfile = st.file_uploader("Upload Questions CSV", type=["csv","xlsx","xls"])
bold_kw = st.checkbox("Bold Keywords in DOCX", value=True)
diff_letter = st.checkbox("Show Difficulty as single letter (L/M/H)", value=False)

if qfile:
    try:
        df = pd.read_csv(qfile)
    except Exception:
        qfile.seek(0)
        df = pd.read_excel(qfile)
    if "Question" not in df.columns:
        st.error("CSV must have a 'Question' column")
    else:
        for col in ["Unit","Subunit","Marks","Answer","Teacher ID","Tag","CO"]:
            if col not in df.columns: df[col] = ""
        df["_Keywords (multi-word)"] = df["Question"].astype(str).apply(
            lambda x: ", ".join(_candidates_rake(x))
        )
        st.subheader("Preview (first 15 rows)")
        st.dataframe(df.head(15))

        if st.button("Generate Question Bank (.docx)"):
            doc = Document()
            cov = doc.add_paragraph()
            cov.add_run(
                f"Question Bank Export\nVersion: {VERSION}\n"
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            ).bold = True
            doc.add_page_break()

            for idx, row in df.iterrows():
                qno = idx + 1
                question = str(row.get("Question",""))
                unit = str(row.get("Unit",""))
                subunit = str(row.get("Subunit",""))
                marks = str(row.get("Marks",""))
                answer = str(row.get("Answer",""))
                teacher = str(row.get("Teacher ID",""))
                tag = str(row.get("Tag",""))
                co = str(row.get("CO","")) or "CO1"

                bloom = detect_bloom_level(question)
                diff = assign_difficulty(bloom)
                qtype = classify_question_type(question)
                keywords_str = ", ".join(_candidates_rake(question)) or "NO_PHRASE_FOUND"

                tbl = doc.add_table(rows=0, cols=2)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.style = "Table Grid"

                _row(tbl, "Question No.", qno)
                _row(tbl, "Question", question)
                _row(tbl, "Unit", f"Unit {unit}")
                _row(tbl, "Subunit", subunit)
                _row(tbl, "Marks", marks)
                _row(tbl, "Difficulty", (diff[0].upper() if diff_letter else diff))
                _row(tbl, "Answer", answer)
                _row(tbl, "Question Type", qtype)
                _row(tbl, "Tag", tag or "[Unit name not found]")
                _row(tbl, "Keywords", keywords_str, bold_val=bold_kw)
                _row(tbl, "Blooms Taxonomy", bloom)
                _row(tbl, "Course Outcome", co)
                _row(tbl, "Teacher ID", f"<{teacher}>")
                _row(tbl, "Year", "<System updates>")
                _row(tbl, "Year asked", "<System updates>")
                _row(tbl, "Frequency", "<System updates>")

                doc.add_page_break()

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button(
                "Download DOCX",
                buf,
                file_name=f"QuestionBank_Output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
